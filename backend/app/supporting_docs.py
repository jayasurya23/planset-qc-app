"""Supporting-document ingestion + evidence layer.

Engineering calculations, CESR surveys, PVSyst reports, ampacity studies and
other external source-of-truth documents are ingested BEFORE the planset is
analyzed. For each document we:

1. Classify it into a known type (``cesr``, ``pvsyst``, ``ampacity``, etc.).
2. Run a type-specific extractor to pull structured specs plus a short
   human-readable summary and excerpt from the source.
3. Emit a ``SupportingDoc`` record that the analyzer threads into the vision
   prompts as an ``Evidence`` block, and runs a dedicated consistency pass
   against the planset values.
"""
from __future__ import annotations

import json
import logging
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------


SUPPORTED_DOC_TYPES = (
    "cesr",                 # Interconnection survey / CESIR / impact study
    "pvsyst",               # PVSyst report
    "ampacity",             # Cable ampacity / sizing study
    "relay",                # Relay / protection coordination study
    "structural",           # Wind / snow / seismic / pile calcs
    "datasheet",            # Generic manufacturer cut sheet
    "module_datasheet",     # PV module datasheet — populates module_voc/vmp/isc/...
    "inverter_datasheet",   # Inverter datasheet — populates inverter_kva/max_vdc/mppt...
    "transformer_datasheet",# Transformer datasheet — populates xfmr_kva/primary_v/...
    "bod",                  # Owner BOD / tech spec / Exhibit / required-docs checklist
    "generic",              # Anything else — still kept as free-text evidence
)


@dataclass
class SupportingDoc:
    filename: str
    doc_type: str
    summary: str = ""
    specs: dict[str, Any] = field(default_factory=dict)
    raw_excerpt: str = ""
    page_count: int = 0
    source_pages: list[int] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


# ---------------------------------------------------------------------------
# Document text/image extraction
# ---------------------------------------------------------------------------


def _read_pdf(data: bytes, max_pages_text: int = 40, max_pages_image: int = 3) -> tuple[str, list[bytes], int]:
    """Return (joined text, first-N page images, page count)."""
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        page_count = doc.page_count
        text_parts: list[str] = []
        for i in range(min(page_count, max_pages_text)):
            text_parts.append(doc[i].get_text("text"))
        pdf_text = "\n---PAGE BREAK---\n".join(text_parts)

        images: list[bytes] = []
        for i in range(min(page_count, max_pages_image)):
            pix = doc[i].get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            images.append(pix.tobytes("png"))
        return pdf_text, images, page_count
    finally:
        doc.close()


def _read_xlsx(data: bytes, max_sheets: int = 8) -> str:
    """Extract text from an Excel workbook (tables → tab-separated text)."""
    try:
        import openpyxl
        from io import BytesIO
    except ImportError:
        return ""
    try:
        wb = openpyxl.load_workbook(BytesIO(data), data_only=True, read_only=True)
    except Exception:
        log.exception("Failed to open xlsx")
        return ""
    parts: list[str] = []
    for sheet_name in wb.sheetnames[:max_sheets]:
        sh = wb[sheet_name]
        parts.append(f"=== SHEET: {sheet_name} ===")
        for row in sh.iter_rows(values_only=True, max_row=500):
            cells = ["" if v is None else str(v) for v in row]
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_docx(data: bytes) -> str:
    """Extract text from a .docx file (paragraphs + table cells, in order)."""
    try:
        from io import BytesIO
        from docx import Document
    except ImportError:
        log.warning("python-docx not installed — docx files will be skipped")
        return ""
    try:
        doc = Document(BytesIO(data))
    except Exception:
        log.exception("Failed to open docx")
        return ""
    parts: list[str] = []
    # Body iteration via the document XML preserves the paragraph/table
    # order; python-docx's Document.paragraphs misses cells inside tables
    # which is where most of an Exhibit V's spec data lives.
    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        parts.append("=== TABLE ===")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(" ".join((p.text or "").strip() for p in cell.paragraphs).strip())
            parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_any(filename: str, data: bytes) -> tuple[str, list[bytes], int]:
    """Return (text, images, page_count) for any file format."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _read_pdf(data)
    if ext in (".xlsx", ".xls"):
        return _read_xlsx(data), [], 0
    if ext == ".docx":
        return _read_docx(data), [], 0
    if ext in (".txt", ".eml", ".msg", ".csv"):
        return data.decode("utf-8", errors="replace")[:80000], [], 0
    # Last resort — try as text
    try:
        return data.decode("utf-8", errors="replace")[:80000], [], 0
    except Exception:
        return "", [], 0


# ---------------------------------------------------------------------------
# Classification
# ---------------------------------------------------------------------------


# Filename hints — order matters: more-specific patterns first so they win
# against the broader generic-utility "interconnect" pattern in cesr.
_FILENAME_HINTS = [
    # BOD-class docs: owner spec / contract exhibit / required-docs checklist
    (re.compile(r"\bexhibit\b|tech(?:nical)?[_\s-]*spec|owner[_\s-]*spec|"
                r"\bbod\b|design[_\s-]*parameters|required[_\s-]*documents",
                re.I), "bod"),
    (re.compile(r"pvsyst|pv[_\s-]*syst|yield.*report|energy.*model", re.I), "pvsyst"),
    (re.compile(r"ampacity|cable.*siz|conductor.*siz|neher", re.I), "ampacity"),
    (re.compile(r"relay|protect.*coord|coord.*study", re.I), "relay"),
    (re.compile(r"struct|wind.*load|snow.*load|seism|pile|foundation", re.I), "structural"),
    (re.compile(r"data.*sheet|cut.*sheet|spec.*sheet|submittal|application[_\s-]*note",
                re.I), "datasheet"),
    # CESR last among the specific patterns — broader trigger
    (re.compile(r"cesr|cesir|sis[_\s-]*report|impact.*study|host.*capacity|"
                r"\biagreement\b|\bia[_\s-]\d|interconnect", re.I), "cesr"),
]


# Content hints — also order-matters; bod content patterns are matched first
# so an Exhibit V tech spec doesn't hit the broad "customer" cesr pattern.
_CONTENT_HINTS = [
    (re.compile(r"\b(basis of design|owner.{0,3}spec|owner.{0,3}criteria|"
                r"technical specification|exhibit [a-z]{1,3}|"
                r"required documents?)\b", re.I), "bod"),
    (re.compile(r"\b(PVsyst|Pnom|specific production|performance ratio|loss diagram)\b",
                re.I), "pvsyst"),
    (re.compile(r"\b(ampacity|NEC 310\.16|NEC 310\.60|ambient adjustment|conduit fill)\b",
                re.I), "ampacity"),
    (re.compile(r"\b(TCC curve|pickup setting|time dial|50/51|relay setting|"
                r"protection coordination)\b", re.I), "relay"),
    (re.compile(r"\b(ASCE|wind speed|snow load|seismic|pile capacity|"
                r"overturning moment|uplift)\b", re.I), "structural"),
    (re.compile(r"\b(datasheet|cut sheet|submittal|mechanical specifications)\b",
                re.I), "datasheet"),
    (re.compile(r"\b(CESR|CESIR|host capacity|impact study|"
                r"interconnection (?:agreement|application))\b", re.I), "cesr"),
]


_DATASHEET_SUBTYPE_HINTS: list[tuple[re.Pattern, str]] = [
    # Module — strongly implied by Voc/Vmp/Isc/Imp + STC ratings, or by
    # known module-vendor model-number patterns.
    (re.compile(
        r"\b(monocrystalline|polycrystalline|bifacial|cell technology|"
        r"PERC|TOPCon|HJT|n-type|p-type)\b|"
        r"\b(Voc|Vmp|Isc|Imp)\b.*\b(STC|standard test conditions)\b|"
        r"\b(JKM\d|JAM\d|CS\d|REC\d|Q\.PEAK|TSM-|LR\d|CHSM|VSUN|"
        r"SunPower|LONGi|Trina|Canadian Solar)\b",
        re.I | re.S),
     "module_datasheet"),
    # Inverter — vendor model numbers, MPPT/kVA AC output language.
    (re.compile(
        r"\b(string inverter|central inverter|MPPT range|max DC input|"
        r"DC input voltage|AC output|nominal AC|AC nominal voltage)\b|"
        r"\b(XGI[ -]?\d|SG\d+|CPS\d+|SMA[- ]|Sungrow|Yaskawa Solectria|"
        r"ABB|Power Electronics|Chint|TBEA)\b",
        re.I),
     "inverter_datasheet"),
    # Transformer — kVA + primary/secondary windings + impedance/BIL.
    (re.compile(
        r"\b(pad[\s-]?mount|step[\s-]?up transformer|primary winding|"
        r"secondary winding|impedance.*%|%Z|BIL|tap changer|delta[\s-]?wye|"
        r"wye[\s-]?delta|HV winding|LV winding)\b",
        re.I),
     "transformer_datasheet"),
]


def _subclassify_datasheet(filename: str, text: str) -> str:
    """If a doc has been classified as 'datasheet', try to narrow to a
    specific equipment type so we can pull the right project_details fields.

    Returns one of:
      module_datasheet | inverter_datasheet | transformer_datasheet | datasheet
    """
    fn = (filename or "").lower()
    if "module" in fn or "panel" in fn:
        return "module_datasheet"
    if "inverter" in fn:
        return "inverter_datasheet"
    if "transformer" in fn or "xfmr" in fn or "xfm" in fn:
        return "transformer_datasheet"

    head = (text or "")[:8000]
    for rx, kind in _DATASHEET_SUBTYPE_HINTS:
        if rx.search(head):
            return kind
    return "datasheet"  # generic fallback


def classify(filename: str, text: str) -> str:
    """Best-effort doc-type classifier. Filename first, then content keywords.

    For documents classified as 'datasheet', a second-pass sub-classifier
    narrows to module/inverter/transformer when possible — this is what
    lets us auto-populate project_details fields downstream.
    """
    matched = None
    for rx, kind in _FILENAME_HINTS:
        if rx.search(filename):
            matched = kind
            break
    if matched is None:
        head = (text or "")[:6000]
        for rx, kind in _CONTENT_HINTS:
            if rx.search(head):
                matched = kind
                break
    if matched is None:
        matched = "generic"
    if matched == "datasheet":
        matched = _subclassify_datasheet(filename, text)
    return matched


# ---------------------------------------------------------------------------
# Extraction prompts  (one per doc type)
# ---------------------------------------------------------------------------


_COMMON_TAIL = """
Return ONLY a JSON object with the keys listed above. Omit keys you cannot
find — do NOT include keys whose value you would have to guess. For each
numeric value include the unit as part of the string (e.g. "8.5 kA", "13.2 kV").
If you found a value on a specific page, include the page number in
"source_pages" (a list of integers). Include a short plain-text one-sentence
"summary" describing what this document is.
"""


_CESR_PROMPT = """\
You are reading a utility interconnection survey / CESIR / impact study for a
solar PV plant. Extract ONLY the authoritative values below — these are the
source of truth against which the planset must be validated.

```json
{
  "doc_type": "cesr",
  "utility_name": "",
  "poi_voltage_ll": "line-to-line POI voltage, e.g. 13.2 kV",
  "poi_voltage_lg": "line-to-ground, e.g. 7.62 kV",
  "fault_current_3ph_ka": "",
  "fault_current_lg_ka": "",
  "x_over_r": "",
  "utility_feeder": "feeder / circuit ID",
  "required_transformer_grounding": "e.g. Wye-grounded-Delta, Wye-grounded-Wye-grounded",
  "grounding_transformer_required": "yes/no + details",
  "required_protection": "list of protection functions required by utility (27, 59, 81, 50/51, 25, DTT, etc.)",
  "direct_transfer_trip_required": "yes/no",
  "anti_islanding_requirement": "",
  "power_factor_requirement": "",
  "voltage_regulation_requirement": "",
  "max_export_kw": "",
  "max_export_kvar": "",
  "metering_accuracy_class": "",
  "special_conditions": "any additional utility-specific requirements",
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_PVSYST_PROMPT = """\
You are reading a PVSyst simulation report for a solar PV plant. Extract the
engineering source-of-truth values below.

```json
{
  "doc_type": "pvsyst",
  "site_location": "",
  "site_lat_long": "",
  "system_type": "fixed-tilt / tracker",
  "dc_capacity_kwp": "",
  "ac_capacity_kva": "",
  "ac_capacity_kw": "",
  "dc_ac_ratio": "",
  "module_make_model": "",
  "module_count": "",
  "inverter_make_model": "",
  "inverter_count": "",
  "is_bifacial": "yes/no",
  "bifaciality_factor": "e.g. 70%",
  "tilt_deg": "",
  "tracker_rotation_range": "e.g. -60 to +60",
  "azimuth_deg": "",
  "pitch_m": "",
  "gcr": "",
  "albedo": "",
  "soiling_loss_pct": "",
  "thermal_uc_ww_m2_k": "e.g. 25",
  "thermal_uv_ww_m3_k": "e.g. 1.2",
  "wiring_loss_pct": "",
  "mismatch_loss_pct": "",
  "lid_loss_pct": "",
  "quality_loss_pct": "",
  "iam_loss_pct": "",
  "near_shading_method": "linear / fast",
  "shading_electrical_effect_pct": "",
  "backtracking_enabled": "yes/no",
  "annual_energy_kwh": "",
  "specific_yield_kwh_kwp": "",
  "performance_ratio_pct": "",
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_AMPACITY_PROMPT = """\
You are reading an electrical cable-sizing / ampacity study. Extract the
derated ampacity results — these are the values the planset cables MUST meet
or exceed.

```json
{
  "doc_type": "ampacity",
  "reference_standard": "e.g. NEC 2020 Table 310.16",
  "ambient_temp_c": "",
  "conductor_temp_rating_c": "e.g. 75 or 90",
  "derating_factors_applied": "e.g. ambient 0.82, bundling 0.80",
  "conduit_fill_max_pct": "",
  "circuits": [
    {
      "circuit_id": "e.g. DC Homerun, INV-01 to SWBD",
      "size": "AWG / kcmil",
      "material": "CU / AL",
      "insulation": "",
      "sets": "number of parallel sets",
      "continuous_load_a": "",
      "derated_ampacity_a": "",
      "source_page": ""
    }
  ],
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_RELAY_PROMPT = """\
You are reading a relay / protection coordination study for a solar PV plant.
Extract the final settings for each relay.

```json
{
  "doc_type": "relay",
  "system_voltage": "",
  "study_standard": "e.g. IEEE 1547-2018, IEEE 242",
  "ct_ratio": "",
  "vt_ratio": "",
  "relays": [
    {
      "location": "e.g. Recloser, Main breaker, 50/51 at POI",
      "functions": "e.g. 27, 59, 81, 50/51, 25, 32",
      "pickup": "",
      "time_delay": "",
      "curve": ""
    }
  ],
  "coordination_interval_s": "",
  "arc_flash_cal_cm2_at_main": "",
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_STRUCTURAL_PROMPT = """\
You are reading structural calculations for a solar PV plant.

```json
{
  "doc_type": "structural",
  "code_reference": "e.g. ASCE 7-16, IBC 2018",
  "wind_speed_mph": "",
  "wind_exposure": "",
  "snow_load_psf": "",
  "seismic_category": "",
  "sds_g": "",
  "sd1_g": "",
  "pile_type": "",
  "pile_embedment_ft": "",
  "module_uplift_psf": "",
  "racking_make_model": "",
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_MODULE_DS_PROMPT = """\
You are reading a PV module datasheet. Extract ONLY the electrical and
thermal characteristics under STC (Standard Test Conditions, 1000 W/m²,
25 °C, AM1.5). These values feed deterministic stringing/fuse calculations
and must match the project's project_details exactly.

Field names below are project_details schema keys — use them verbatim so
they auto-merge.

```json
{
  "doc_type": "module_datasheet",
  "module_make": "manufacturer name",
  "module_model": "model / part number — exactly as printed on the datasheet",
  "module_stc_watts": "STC nameplate Pmax in W (numeric only, no unit)",
  "module_voc": "open-circuit voltage at STC in V (numeric)",
  "module_vmp": "max-power voltage at STC in V (numeric)",
  "module_isc": "short-circuit current at STC in A (numeric)",
  "module_imp": "max-power current at STC in A (numeric)",
  "module_temp_coeff_voc": "TCoeff_Voc in %/°C — typically negative, e.g. -0.27",
  "module_temp_coeff_isc": "TCoeff_Isc in %/°C",
  "module_temp_coeff_pmax": "TCoeff_Pmax in %/°C",
  "is_bifacial": "yes/no",
  "summary": "one sentence describing the module",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_INVERTER_DS_PROMPT = """\
You are reading an inverter datasheet (string or central inverter).
Extract the electrical specs that drive sizing/stringing/grounding rules.

Field names below are project_details schema keys — use them verbatim.

```json
{
  "doc_type": "inverter_datasheet",
  "inverter_make": "manufacturer",
  "inverter_model": "exact model / part number",
  "inverter_kva": "AC nameplate kVA (numeric, no unit)",
  "inverter_kw": "AC nameplate kW (numeric)",
  "inverter_max_vdc": "max DC input voltage in V (numeric — used for cold-Voc check)",
  "inverter_mppt_range": "MPPT range as 'min-max V' string, e.g. '200-1000V'",
  "inverter_max_dc_input_isc": "max DC input current per MPPT in A (numeric)",
  "inverter_ac_voltage": "AC nominal voltage in V (numeric)",
  "inverter_max_short_circuit_amps": "AC short-circuit contribution in A (numeric)",
  "inverter_efficiency_pct": "max CEC or Euro efficiency in %",
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_TRANSFORMER_DS_PROMPT = """\
You are reading a transformer datasheet / shop drawing — typically the
medium-voltage step-up between inverter AC bus and the utility POI.
Extract specs that drive sizing, BIL, and grounding rules.

Field names below are project_details schema keys.

```json
{
  "doc_type": "transformer_datasheet",
  "transformer_kva": "per-unit nameplate kVA (numeric — what's printed on this datasheet)",
  "transformer_quantity": "count of step-ups in this project. Often shown as '(2) 2500 kVA'. Defaults to 1 if not stated.",
  "transformer_primary_voltage": "primary L-L voltage in V (numeric, e.g. 13200)",
  "transformer_secondary_voltage": "secondary L-L voltage in V (numeric, e.g. 600)",
  "transformer_winding_config": "e.g. 'Delta-Wye-grounded', 'Wye-grounded-Wye-grounded'",
  "transformer_impedance": "%Z, e.g. 5.75",
  "transformer_bil": "BIL in kV, e.g. 95",
  "transformer_cooling_class": "ONAN / ONAN-ONAF / etc.",
  "transformer_taps": "tap changer description, e.g. '+/-2 x 2.5%'",
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_BOD_PROMPT = """\
You are reading a project's Basis of Design (BOD) / owner technical
specification / contract Exhibit / required-documents checklist for a
solar PV plant. Extract ONLY the authoritative project-specific
parameters and procurement/spec requirements that the planset must
conform to. These are the values planset rules tagged "per project
spec / owner criteria / per BOD" depend on.

Pull whatever is present — fields with no information should be omitted.

```json
{
  "doc_type": "bod",
  "owner_name": "",
  "epc_name": "",
  "project_size_mwac": "",
  "project_size_mwdc": "",
  "module_make_model": "approved or required module",
  "inverter_make_model": "",
  "transformer_kva_each": "",
  "transformer_winding_config": "",
  "fence_spec": "type/height/post material/gate width — e.g. '7ft fixed-knot, metal posts, 12ft swing gates'",
  "fence_grounding_threshold": "ft from array, per SOP",
  "ground_ring_burial_depth": "",
  "grounding_material": "copper/CCS/steel — per geotech",
  "access_road_width_tier1_ft": "",
  "access_road_width_tier2_ft": "",
  "fire_access_width_ft": "",
  "laydown_size_tier": "",
  "met_station_count_formula": "e.g. '2 min + 1 per 20MWdc above 40MWdc'",
  "monitoring_backup_hours": "",
  "voltage_drop_max_pct": "",
  "design_temp_low_c": "",
  "design_temp_high_c": "",
  "ambient_temp_c": "",
  "wind_speed_mph": "",
  "snow_load_psf": "",
  "racking_make_model": "",
  "racking_type": "fixed/tracker",
  "tracker_rotation_range_deg": "",
  "approved_equipment_list": "list — e.g. 'XGI-1500, Generac genset, ABB SCADA'",
  "required_studies": "e.g. 'arc flash study, short-circuit study, ground grid study'",
  "deliverable_stages": "list — e.g. ['30%','60%','90%','IFC','As-Built']",
  "special_conditions": "non-standard requirements unique to this project",
  "summary": "one sentence describing what this BOD covers",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_GENERIC_PROMPT = """\
Summarize this engineering document. Return ONLY JSON:

```json
{
  "doc_type": "generic",
  "detected_subject": "short — e.g. 'Transformer datasheet', 'Interconnection email'",
  "key_values": { "label": "value", "...": "..." },
  "summary": "one sentence",
  "source_pages": []
}
```
""" + _COMMON_TAIL


_PROMPTS = {
    "cesr": _CESR_PROMPT,
    "pvsyst": _PVSYST_PROMPT,
    "ampacity": _AMPACITY_PROMPT,
    "relay": _RELAY_PROMPT,
    "structural": _STRUCTURAL_PROMPT,
    "datasheet": _GENERIC_PROMPT,
    "module_datasheet": _MODULE_DS_PROMPT,
    "inverter_datasheet": _INVERTER_DS_PROMPT,
    "transformer_datasheet": _TRANSFORMER_DS_PROMPT,
    "bod": _BOD_PROMPT,
    "generic": _GENERIC_PROMPT,
}


# Datasheet sub-types whose extracted specs use project_details schema keys
# directly. Their specs can be merged into the analysis call's
# ``project_details`` so calc rules (validate_stringing, validate_fuse_sizing,
# etc.) stop deferring on "Missing module_voc" / "Missing inverter_kva".
_AUTO_MERGE_DATASHEET_TYPES = {
    "module_datasheet",
    "inverter_datasheet",
    "transformer_datasheet",
}


def project_details_from_supporting_docs(
    supporting_docs: list[dict] | None,
) -> dict[str, str]:
    """Aggregate auto-extractable project_details fields from uploaded
    datasheets. The caller (typically /api/analyze) merges this with the
    user's own project_details — user-entered values WIN to preserve any
    manual overrides.
    """
    if not supporting_docs:
        return {}
    merged: dict[str, str] = {}
    for d in supporting_docs:
        if d.get("doc_type") not in _AUTO_MERGE_DATASHEET_TYPES:
            continue
        specs = d.get("specs") or {}
        for k, v in specs.items():
            if v is None:
                continue
            sv = str(v).strip()
            if not sv:
                continue
            # First non-empty value wins. If a user uploads two module
            # datasheets, the first one's values stand — keep it
            # deterministic and let the user resolve conflicts manually.
            if k not in merged:
                merged[k] = sv
    return merged


# ---------------------------------------------------------------------------
# Extraction runner
# ---------------------------------------------------------------------------


def _extract_json(text: str) -> dict:
    """Best-effort JSON-object extraction (same style as gemini_analyzer)."""
    m = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass
    try:
        parsed = json.loads(text)
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        pass
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end > start:
        try:
            return json.loads(text[start:end + 1])
        except json.JSONDecodeError:
            pass
    return {}


def process_document(
    filename: str,
    data: bytes,
    *,
    deep: bool = False,
) -> SupportingDoc:
    """Ingest a single supporting document end-to-end."""
    from .gemini_client import analyze_multiple_images, analyze_text

    text, images, page_count = _read_any(filename, data)
    doc_type = classify(filename, text)
    prompt_body = _PROMPTS.get(doc_type, _GENERIC_PROMPT)

    # Build the full prompt: file context + schema + instructions.
    header = (
        f"File name: {filename}\n"
        f"File type: {Path(filename).suffix.lower() or 'text'}\n"
        f"Detected document type: {doc_type}\n\n"
    )
    body = header
    if text:
        body += f"=== EXTRACTED TEXT ===\n{text[:60000]}\n=== END TEXT ===\n\n"
    body += prompt_body

    try:
        if images:
            raw = analyze_multiple_images(images, body, "image/png", deep=deep)
        else:
            raw = analyze_text(body, deep=deep)
    except Exception:
        log.exception("Supporting-doc AI call failed for %s", filename)
        raw = ""

    parsed = _extract_json(raw)
    specs: dict[str, Any] = {}
    summary = ""
    source_pages: list[int] = []
    for k, v in parsed.items():
        if k == "summary" and isinstance(v, str):
            summary = v.strip()
        elif k == "source_pages" and isinstance(v, list):
            source_pages = [int(x) for x in v if str(x).isdigit()]
        elif k == "doc_type":
            continue
        else:
            if v is None:
                continue
            if isinstance(v, str) and not v.strip():
                continue
            specs[k] = v

    # A short raw excerpt helps later retrieval / audit trails.
    excerpt = (text or "")[:4000]

    return SupportingDoc(
        filename=filename,
        doc_type=doc_type,
        summary=summary,
        specs=specs,
        raw_excerpt=excerpt,
        page_count=page_count,
        source_pages=source_pages,
    )


# ---------------------------------------------------------------------------
# Evidence-context rendering (goes into planset check prompts)
# ---------------------------------------------------------------------------


def _fmt_kv(key: str, value: Any) -> str:
    label = key.replace("_", " ")
    if isinstance(value, (list, dict)):
        s = json.dumps(value, ensure_ascii=False, separators=(", ", ": "))
        if len(s) > 180:
            s = s[:177] + "..."
        return f"    {label}: {s}"
    return f"    {label}: {value}"


def build_evidence_context(docs: list[dict] | None) -> str:
    """Render the evidence block that appends to every vision-check prompt."""
    if not docs:
        return ""

    lines: list[str] = [
        "",
        "=== SUPPORTING DOCUMENTS (engineering source of truth) ===",
        "The following values come from external engineering calculations",
        "and studies for this project. They are AUTHORITATIVE. When you",
        "audit the planset, COMPARE values on the drawing against these",
        "and flag any mismatch as a Fail with both values quoted",
        "(format: 'Planset: X / <DocType>: Y').",
        "",
    ]
    for idx, d in enumerate(docs, start=1):
        filename = d.get("filename", "?")
        doc_type = d.get("doc_type", "generic")
        summary = d.get("summary", "").strip()
        specs = d.get("specs") or {}
        src_pages = d.get("source_pages") or []
        pg_str = f" (pp. {', '.join(str(p) for p in src_pages)})" if src_pages else ""
        lines.append(f"[{idx}] {filename}  —  type: {doc_type}{pg_str}")
        if summary:
            lines.append(f"    summary: {summary}")
        if specs:
            for k, v in specs.items():
                lines.append(_fmt_kv(k, v))
        lines.append("")
    lines.append("=== END SUPPORTING DOCUMENTS ===")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Cross-check prompt builder (for the dedicated consistency category)
# ---------------------------------------------------------------------------


CONSISTENCY_CHECK_CATEGORY = "Supporting Document Consistency"


CONSISTENCY_PROMPT = """\
You are a senior electrical PE auditing a solar PV planset against the
engineering source-of-truth documents (CESIR survey, PVSyst report, ampacity
study, etc.) provided above.

Your job: for EACH material spec in the supporting documents, READ the
corresponding value on the planset pages shown below and report whether they
match.

For EVERY comparison emit ONE JSON finding:
- "check": short snake_case name (e.g. "planset_vs_cesr_fault_current").
- "status": "Pass" | "Fail" | "Needs Review".
- "value": the value you read on the planset.
- "evidence": one sentence showing both values, like
    "Planset shows 8kA AIC; CESIR (p.3) requires 8.5kA → under-rated."
- "location": table/row/sheet where the planset value lives.
- "location_text": a literal short text excerpt searchable on the planset page.
- "source_doc": the filename of the supporting document you compared against.
- "severity": "low" | "medium" | "high" (mismatches in ratings are "high").

Rules:
- If the planset lacks a value that the supporting doc specifies, mark
  "Needs Review" with evidence "Not found on provided planset pages."
- If the supporting-doc value is ambiguous / absent, SKIP it (do not emit).
- Tolerances: voltages, currents, and ratings must match within 2%.
  Percentages (loss stacks, PR) match within 0.5 absolute points.

Only return a JSON array of findings. No other text.
"""
